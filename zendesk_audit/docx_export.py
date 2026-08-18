"""
Convert a Zendesk article's HTML body into a .docx file.

Heading mapping is the reverse of the QRG migration tool's style map so
round-tripping through both tools preserves structure:
    h2 -> Heading 1,  h3 -> Heading 2,  h4 -> Heading 3
"""

import io
import re
from urllib.parse import urlparse

import requests
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

import audit


HEADING_MAP = {"h2": 1, "h3": 2, "h4": 3}


def fetch_article(article_id):
    resp = audit.zendesk_request(
        "get", f"/api/v2/help_center/articles/{article_id}.json"
    )
    return resp.json().get("article", {})


def _is_zendesk_url(url):
    if not audit.ZENDESK_SUBDOMAIN:
        return False
    return audit.ZENDESK_SUBDOMAIN in urlparse(url).hostname


def _download_image(url):
    try:
        kwargs = {"timeout": 15}
        if _is_zendesk_url(url):
            kwargs["auth"] = audit.zendesk_auth()
        r = requests.get(url, **kwargs)
        if r.status_code == 200 and r.headers.get("content-type", "").startswith("image"):
            return io.BytesIO(r.content)
    except Exception:
        pass
    return None


def _add_image(doc, src, alt=""):
    img_stream = _download_image(src) if src.startswith("http") else None
    if img_stream:
        try:
            doc.add_picture(img_stream, width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return True
        except Exception:
            pass
    p = doc.add_paragraph()
    p.add_run(f"[Image: {alt or src}]")
    p.style = doc.styles["Normal"]
    return False


def _add_inline_runs(paragraph, element, doc=None):
    if isinstance(element, NavigableString):
        text = str(element)
        if text.strip():
            paragraph.add_run(text)
        return

    if not isinstance(element, Tag):
        return

    tag = element.name

    if tag in ("strong", "b"):
        run = paragraph.add_run(element.get_text())
        run.bold = True
    elif tag in ("em", "i"):
        run = paragraph.add_run(element.get_text())
        run.italic = True
    elif tag == "a":
        run = paragraph.add_run(element.get_text())
        run.underline = True
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    elif tag == "code":
        run = paragraph.add_run(element.get_text())
        run.font.name = "Consolas"
        run.font.size = Pt(9)
    elif tag == "br":
        paragraph.add_run("\n")
    elif tag == "img":
        if doc:
            _add_image(doc, element.get("src", ""), element.get("alt", ""))
        else:
            alt = element.get("alt", element.get("src", ""))
            paragraph.add_run(f"[Image: {alt}]")
    else:
        for child in element.children:
            _add_inline_runs(paragraph, child, doc=doc)


def _process_list(doc, list_tag, indent=0):
    ordered = list_tag.name == "ol"
    counter = 1
    for li in list_tag.find_all("li", recursive=False):
        prefix = f"{counter}. " if ordered else "• "
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5 * (indent + 1))
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.add_run(prefix)

        for child in li.children:
            if isinstance(child, Tag) and child.name in ("ul", "ol"):
                _process_list(doc, child, indent + 1)
            elif isinstance(child, Tag) and child.name == "img":
                _add_image(doc, child.get("src", ""), child.get("alt", ""))
            else:
                _add_inline_runs(p, child, doc=doc)

        counter += 1


def _process_table(doc, table_tag):
    rows = table_tag.find_all("tr")
    if not rows:
        return

    col_count = max(
        len(row.find_all(["td", "th"])) for row in rows
    )
    if col_count == 0:
        return

    tbl = doc.add_table(rows=len(rows), cols=col_count)
    tbl.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        cells = row.find_all(["td", "th"])
        for c_idx, cell in enumerate(cells):
            if c_idx < col_count:
                tbl_cell = tbl.cell(r_idx, c_idx)
                tbl_cell.text = ""
                p = tbl_cell.paragraphs[0]
                for child in cell.children:
                    _add_inline_runs(p, child)
                if cell.name == "th":
                    for run in p.runs:
                        run.bold = True


def _process_element(doc, element):
    if isinstance(element, NavigableString):
        text = str(element).strip()
        if text:
            doc.add_paragraph(text)
        return

    if not isinstance(element, Tag):
        return

    tag = element.name

    if tag in HEADING_MAP:
        doc.add_heading(element.get_text(strip=True), level=HEADING_MAP[tag])

    elif tag == "h1":
        doc.add_heading(element.get_text(strip=True), level=1)

    elif tag == "img":
        _add_image(doc, element.get("src", ""), element.get("alt", ""))

    elif tag == "p":
        imgs = element.find_all("img")
        non_img_content = element.get_text(strip=True)
        if imgs and not non_img_content:
            for img in imgs:
                _add_image(doc, img.get("src", ""), img.get("alt", ""))
        else:
            p = doc.add_paragraph()
            for child in element.children:
                _add_inline_runs(p, child, doc=doc)

    elif tag in ("ul", "ol"):
        _process_list(doc, element)

    elif tag == "table":
        _process_table(doc, element)

    elif tag == "div":
        cls = element.get("class", [])
        if "callout" in cls:
            p = doc.add_paragraph()
            p.style = doc.styles["Normal"]
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run("⚠ ")
            run.bold = True
            for child in element.children:
                _add_inline_runs(p, child, doc=doc)
        else:
            for child in element.children:
                _process_element(doc, child)

    elif tag == "pre":
        p = doc.add_paragraph()
        run = p.add_run(element.get_text())
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    elif tag == "hr":
        p = doc.add_paragraph()
        p.add_run("_" * 60)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    elif tag == "blockquote":
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.add_run("“")
        for child in element.children:
            _add_inline_runs(p, child, doc=doc)
        p.add_run("”")

    elif tag == "figure":
        img = element.find("img")
        if img:
            _add_image(doc, img.get("src", ""), img.get("alt", ""))
        caption = element.find("figcaption")
        if caption:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(caption.get_text(strip=True))
            run.italic = True
            run.font.size = Pt(9)


def html_to_docx(title, html_body):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    heading = doc.add_heading(title, level=1)
    heading.style = doc.styles["Heading 1"]

    if not html_body:
        return doc

    soup = BeautifulSoup(html_body, "html.parser")

    for element in soup.children:
        _process_element(doc, element)

    return doc


def export_article_docx(article_id):
    article = fetch_article(article_id)
    if not article:
        return None, None

    title = article.get("title", "Untitled")
    html_body = article.get("body", "")

    doc = html_to_docx(title, html_body)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_title = re.sub(r'[^\w\s-]', '', title).strip().replace(' ', '_')[:80]
    filename = f"{safe_title}.docx"

    return buf, filename
